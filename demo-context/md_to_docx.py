"""Convert demo-questions-review.md into three per-variant Word documents."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/thuannguyen/Documents/Aperia/code/visual-to-code-poc/demo-context"
MD   = f"{BASE}/demo-questions-review.md"

VARIANTS = {
    "Clover Variant":        f"{BASE}/demo-questions-clover.docx",
    "Business Owner Variant": f"{BASE}/demo-questions-business-owner.docx",
    "ISO Variant":            f"{BASE}/demo-questions-iso.docx",
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_bold_inline(para, text):
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def parse_table(lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-: |]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def apply_table_style(table):
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if not para.runs:
                para.clear()
                run = para.add_run(para.text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for idx, row in enumerate(table.rows[1:], 1):
        color = "D9E1F2" if idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            set_cell_bg(cell, color)


def flush_table(doc, table_buffer):
    if not table_buffer:
        return
    rows = parse_table(table_buffer)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = tbl.cell(r_idx, c_idx)
            cell.paragraphs[0].clear()
            add_bold_inline(cell.paragraphs[0], cell_text)
    apply_table_style(tbl)
    doc.add_paragraph()


def render_lines(doc, lines):
    """Render a list of markdown lines into doc."""
    i = 0
    table_buffer = []

    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        if stripped.startswith("|"):
            table_buffer.append(stripped)
            i += 1
            continue

        if table_buffer:
            flush_table(doc, table_buffer)
            table_buffer = []

        if not stripped:
            i += 1
            continue

        if re.match(r"^---+$", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            p = doc.add_paragraph(style=heading_map.get(level, "Heading 4"))
            add_bold_inline(p, text)
            i += 1
            continue

        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            if text:
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(text)
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_bold_inline(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph(style="Normal")
        add_bold_inline(p, stripped)
        i += 1

    flush_table(doc, table_buffer)


def make_doc(lines, out_path):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    render_lines(doc, lines)
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ── Split markdown by H2 variant heading ─────────────────────────────────────

with open(MD, encoding="utf-8") as f:
    all_lines = f.readlines()

# Find the line index of each variant H2 heading
variant_starts = {}
for idx, line in enumerate(all_lines):
    for variant_key in VARIANTS:
        # e.g. "## Clover Variant (`?embed=clover`)"
        if re.match(rf"^##\s+{re.escape(variant_key)}", line.strip()):
            variant_starts[variant_key] = idx

sorted_keys = sorted(variant_starts, key=lambda k: variant_starts[k])

def extract_section(key):
    start = variant_starts[key]
    ki = sorted_keys.index(key)
    end = variant_starts[sorted_keys[ki + 1]] if ki + 1 < len(sorted_keys) else len(all_lines)
    lines = []
    for line in all_lines[start:end]:
        m = re.match(r"^(#{2,})(.*)", line)
        if m:
            line = "#" * (len(m.group(1)) - 1) + m.group(2) + "\n"
        lines.append(line)
    return lines


for key in sorted_keys:
    section_lines = extract_section(key)
    make_doc(section_lines, VARIANTS[key])
